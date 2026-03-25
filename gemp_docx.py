import os
import tempfile
from typing import Any, Dict, List
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    Image as RLImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GEMP_LOGO_PATH = os.getenv(
    "GEMP_LOGO_PATH",
    os.path.join(BASE_DIR, "assets", "energy.jpg"),
)

MONTH_ORDER = [
    "January",
    "February",
    "March",
    "April",
    "May",
    "June",
    "July",
    "August",
    "September",
    "October",
    "November",
    "December",
]


def fmt(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    return "" if s == "-" else s


def _get_first(src: Dict[str, Any], *keys: str) -> Any:
    for key in keys:
        if key in src and src.get(key) is not None:
            value = src.get(key)
            if str(value).strip() != "":
                return value
    return ""


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def set_paragraph_text(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_underlined_value(cell, label: str, value: str):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    label_run = p.add_run(label)
    label_run.bold = True
    label_run.font.name = "Arial"
    label_run.font.size = Pt(10)

    value_run = p.add_run(f" {value or ' '}")
    value_run.font.name = "Arial"
    value_run.font.size = Pt(10)

    set_cell_border(
        cell,
        bottom={"val": "single", "sz": "8", "space": "0", "color": "808080"},
    )


def set_col_width(cell, inches: float):
    cell.width = Inches(inches)


def build_month_rows(rows: List[Dict[str, Any]], stats: Dict[str, Any]) -> List[Dict[str, Any]]:
    row_map = {}
    for r in rows:
        month_name = str(r.get("month", "")).strip()
        if month_name:
            row_map[month_name] = r

    final_rows = []
    for month in MONTH_ORDER:
        src = row_map.get(month, {})
        final_rows.append(
            {
                "month": month,
                "baseline2025": fmt(_get_first(src, "baseline2025", "baseline2016")),
                "buildingDesc": fmt(_get_first(src, "buildingDesc", "buildingDescription")),
                "grossArea": fmt(src.get("grossArea")),
                "airconArea": fmt(src.get("airconArea")),
                "occupants": fmt(src.get("occupants")),
                "kwh": fmt(src.get("kwh")),
            }
        )

    final_rows.append(
        {
            "month": "Average",
            "baseline2025": fmt(stats.get("avgBaseline")),
            "buildingDesc": "",
            "grossArea": fmt(stats.get("avgGrossArea")),
            "airconArea": fmt(stats.get("avgAirconArea")),
            "occupants": fmt(stats.get("avgOccupants")),
            "kwh": fmt(stats.get("avgKwh")),
        }
    )

    return final_rows


def build_gemp_docx(payload: Dict[str, Any]) -> str:
    header = payload.get("header", {}) or {}
    rows = payload.get("rows", []) or []
    stats = payload.get("stats", {}) or {}

    year = fmt(header.get("year")) or "2020"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    p_annex = doc.add_paragraph()
    p_annex.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_annex = p_annex.add_run('"ANNEX A"')
    run_annex.bold = True
    run_annex.font.name = "Arial"
    run_annex.font.size = Pt(13)

    if os.path.exists(GEMP_LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(GEMP_LOGO_PATH, width=Inches(0.8))

    p = doc.add_paragraph()
    set_paragraph_text(
        p,
        "DEPARTMENT OF ENERGY",
        bold=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    p = doc.add_paragraph()
    set_paragraph_text(
        p,
        "Energy Center, Rizal Drive, Bonifacio Global City, Taguig City",
        bold=True,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    p = doc.add_paragraph()
    set_paragraph_text(
        p,
        "Telefax: (632) 8840-2243,  Email: doe.gemp@gmail.com",
        bold=True,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    set_paragraph_text(
        p,
        "GOVERNMENT ENERGY MANAGEMENT PROGRAM",
        bold=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    p = doc.add_paragraph()
    set_paragraph_text(
        p,
        f"Monthly Electricity Consumption Report, {year}",
        bold=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_paragraph()

    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False

    left_values = [
        ("Agency:", fmt(header.get("agency"))),
        ("Address:", fmt(header.get("address"))),
        ("Region:", fmt(header.get("region"))),
    ]
    right_values = [
        ("Tel. Nos.:", fmt(header.get("tel"))),
        ("Fax Nos.:", fmt(header.get("fax"))),
        ("", ""),
    ]

    for i in range(3):
        left_cell = info_table.cell(i, 0)
        right_cell = info_table.cell(i, 1)

        set_col_width(left_cell, 3.75)
        set_col_width(right_cell, 3.75)

        add_underlined_value(left_cell, left_values[i][0], left_values[i][1])
        add_underlined_value(right_cell, right_values[i][0], right_values[i][1])

    doc.add_paragraph()

    table = doc.add_table(rows=2, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [0.9, 1.0, 1.1, 1.05, 1.2, 0.95, 1.25]

    letters = ["A", "B", "C", "D", "E", "F", "G"]
    headers = [
        "Month",
        "Monthly\nConsumption\nBaseline, 2015",
        "Building\nDescription",
        "Gross  Area\n(Square meters)",
        "Air- Conditioned\nArea\n(square meters)",
        "Number of\nOccupants",
        "Monthly Consumption,\nkWh",
    ]

    for col_idx in range(7):
        top_cell = table.cell(0, col_idx)
        hdr_cell = table.cell(1, col_idx)

        set_col_width(top_cell, col_widths[col_idx])
        set_col_width(hdr_cell, col_widths[col_idx])

        set_cell_text(top_cell, letters[col_idx], bold=True, size=8)
        set_cell_text(hdr_cell, headers[col_idx], bold=False, size=9)

        shade_cell(hdr_cell, "F2F2F2")

        set_cell_border(
            top_cell,
            top={"val": "single", "sz": "10", "space": "0", "color": "000000"},
            left={"val": "single", "sz": "10", "space": "0", "color": "000000"},
            bottom={"val": "single", "sz": "10", "space": "0", "color": "000000"},
            right={"val": "single", "sz": "10", "space": "0", "color": "000000"},
        )
        set_cell_border(
            hdr_cell,
            top={"val": "single", "sz": "10", "space": "0", "color": "000000"},
            left={"val": "single", "sz": "10", "space": "0", "color": "000000"},
            bottom={"val": "single", "sz": "10", "space": "0", "color": "000000"},
            right={"val": "single", "sz": "10", "space": "0", "color": "000000"},
        )

    full_rows = build_month_rows(rows, stats)

    for r in full_rows:
        row_cells = table.add_row().cells
        values = [
            fmt(r.get("month")),
            fmt(r.get("baseline2025")),
            fmt(r.get("buildingDesc")),
            fmt(r.get("grossArea")),
            fmt(r.get("airconArea")),
            fmt(r.get("occupants")),
            fmt(r.get("kwh")),
        ]

        for idx, value in enumerate(values):
            set_col_width(row_cells[idx], col_widths[idx])
            set_cell_text(
                row_cells[idx],
                value,
                bold=(fmt(r.get("month")) == "Average"),
                size=9,
                align=WD_ALIGN_PARAGRAPH.LEFT if idx in [0, 2] else WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_border(
                row_cells[idx],
                top={"val": "single", "sz": "8", "space": "0", "color": "000000"},
                left={"val": "single", "sz": "8", "space": "0", "color": "000000"},
                bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
                right={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            )

    doc.add_paragraph()
    doc.add_paragraph()

    prepared_by = fmt(header.get("preparedBy"))
    prepared_by_designation = fmt(header.get("preparedByDesignation")) or "Designation"
    noted_by = fmt(header.get("notedBy"))
    noted_by_designation = fmt(header.get("notedByDesignation")) or "Designation"

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False

    left = sign_table.cell(0, 0)
    right = sign_table.cell(0, 1)

    set_col_width(left, 3.6)
    set_col_width(right, 3.6)

    set_cell_text(left, "Prepared by:", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(right, "Noted by:", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    left2 = sign_table.cell(1, 0)
    right2 = sign_table.cell(1, 1)

    left2.text = ""
    p1 = left2.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(
        f"\n\n{prepared_by if prepared_by else '_______________________________'}\n{prepared_by_designation}"
    )
    r1.font.name = "Arial"
    r1.font.size = Pt(10)

    right2.text = ""
    p2 = right2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(
        f"\n\n{noted_by if noted_by else '_______________________________'}\n{noted_by_designation}"
    )
    r2.font.name = "Arial"
    r2.font.size = Pt(10)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


def _p_style(name: str, size: int = 9, bold: bool = False, align: int = TA_LEFT):
    return ParagraphStyle(
        name=name,
        fontName="Helvetica-Bold" if bold else "Helvetica",
        fontSize=size,
        leading=size + 1,
        alignment=align,
        spaceBefore=0,
        spaceAfter=0,
    )


def build_gemp_pdf(payload: Dict[str, Any]) -> str:
    header = payload.get("header", {}) or {}
    rows = payload.get("rows", []) or []
    stats = payload.get("stats", {}) or {}
    year = fmt(header.get("year")) or "2020"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp.close()

    doc = SimpleDocTemplate(
        tmp.name,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=18,
        bottomMargin=24,
    )

    annex_style = _p_style("annex", size=12, bold=True, align=TA_RIGHT)
    center_12_bold = _p_style("center12bold", size=12, bold=True, align=TA_CENTER)
    center_10_bold = _p_style("center10bold", size=10, bold=True, align=TA_CENTER)
    center_11_bold = _p_style("center11bold", size=11, bold=True, align=TA_CENTER)
    info_style = _p_style("info", size=9, bold=False, align=TA_LEFT)
    table_hdr = _p_style("table_hdr", size=7, bold=False, align=TA_CENTER)
    table_letter = _p_style("table_letter", size=7, bold=True, align=TA_CENTER)
    table_left = _p_style("table_left", size=7, bold=False, align=TA_LEFT)
    table_center = _p_style("table_center", size=7, bold=False, align=TA_CENTER)
    table_avg_left = _p_style("table_avg_left", size=7, bold=True, align=TA_LEFT)
    table_avg_center = _p_style("table_avg_center", size=7, bold=True, align=TA_CENTER)
    sign_lbl = _p_style("sign_lbl", size=9, bold=True, align=TA_LEFT)
    sign_name = _p_style("sign_name", size=9, bold=False, align=TA_LEFT)
    sign_desig = _p_style("sign_desig", size=9, bold=False, align=TA_LEFT)

    story = []

    story.append(Paragraph('"ANNEX A"', annex_style))
    story.append(Spacer(1, 4))

    if os.path.exists(GEMP_LOGO_PATH):
        try:
            logo = RLImage(GEMP_LOGO_PATH, width=44, height=44)
            logo.hAlign = "CENTER"
            story.append(logo)
            story.append(Spacer(1, 6))
        except Exception:
            pass

    story.append(Paragraph("DEPARTMENT OF ENERGY", center_12_bold))
    story.append(
        Paragraph(
            "Energy Center, Rizal Drive, Bonifacio Global City, Taguig City",
            center_10_bold,
        )
    )
    story.append(
        Paragraph(
            "Telefax: (632) 8840-2243,  Email: doe.gemp@gmail.com",
            center_10_bold,
        )
    )
    story.append(Spacer(1, 10))
    story.append(Paragraph("GOVERNMENT ENERGY MANAGEMENT PROGRAM", center_12_bold))
    story.append(Paragraph(f"Monthly Electricity Consumption Report, {year}", center_11_bold))
    story.append(Spacer(1, 10))

    info_data = [
        [
            Paragraph(f"<b>Agency:</b> {xml_escape(fmt(header.get('agency')))}", info_style),
            Paragraph(f"<b>Tel. Nos.:</b> {xml_escape(fmt(header.get('tel')))}", info_style),
        ],
        [
            Paragraph(f"<b>Address:</b> {xml_escape(fmt(header.get('address')))}", info_style),
            Paragraph(f"<b>Fax Nos.:</b> {xml_escape(fmt(header.get('fax')))}", info_style),
        ],
        [
            Paragraph(f"<b>Region:</b> {xml_escape(fmt(header.get('region')))}", info_style),
            Paragraph("", info_style),
        ],
    ]

    info_table = Table(info_data, colWidths=[doc.width / 2.0, doc.width / 2.0])
    info_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(info_table)
    story.append(Spacer(1, 14))

    full_rows = build_month_rows(rows, stats)

    col_widths = [58, 84, 82, 76, 92, 74, 73]

    table_data = [
        [
            Paragraph("A", table_letter),
            Paragraph("B", table_letter),
            Paragraph("C", table_letter),
            Paragraph("D", table_letter),
            Paragraph("E", table_letter),
            Paragraph("F", table_letter),
            Paragraph("G", table_letter),
        ],
        [
            Paragraph("Month", table_hdr),
            Paragraph("Monthly<br/>Consumption<br/>Baseline, 2015", table_hdr),
            Paragraph("Building<br/>Description", table_hdr),
            Paragraph("Gross Area<br/>(Square meters)", table_hdr),
            Paragraph("Air-Conditioned<br/>Area<br/>(square meters)", table_hdr),
            Paragraph("Number of<br/>Occupants", table_hdr),
            Paragraph("Monthly Consumption,<br/>kWh", table_hdr),
        ],
    ]

    for row in full_rows:
        is_avg = fmt(row.get("month")) == "Average"
        left_style = table_avg_left if is_avg else table_left
        center_style = table_avg_center if is_avg else table_center

        table_data.append(
            [
                Paragraph(xml_escape(fmt(row.get("month"))), left_style),
                Paragraph(xml_escape(fmt(row.get("baseline2025"))), center_style),
                Paragraph(xml_escape(fmt(row.get("buildingDesc"))), left_style),
                Paragraph(xml_escape(fmt(row.get("grossArea"))), center_style),
                Paragraph(xml_escape(fmt(row.get("airconArea"))), center_style),
                Paragraph(xml_escape(fmt(row.get("occupants"))), center_style),
                Paragraph(xml_escape(fmt(row.get("kwh"))), center_style),
            ]
        )

    main_table = Table(table_data, colWidths=col_widths, repeatRows=2)
    main_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F2F2F2")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, 1), "CENTER"),
                ("ALIGN", (0, 2), (0, -1), "LEFT"),
                ("ALIGN", (2, 2), (2, -1), "LEFT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(main_table)
    story.append(Spacer(1, 22))

    prepared_by = fmt(header.get("preparedBy"))
    prepared_by_designation = fmt(header.get("preparedByDesignation")) or "Designation"
    noted_by = fmt(header.get("notedBy"))
    noted_by_designation = fmt(header.get("notedByDesignation")) or "Designation"

    sign_table = Table(
        [
            [Paragraph("Prepared by:", sign_lbl), Paragraph("Noted by:", sign_lbl)],
            [
                Paragraph(xml_escape(prepared_by if prepared_by else " "), sign_name),
                Paragraph(xml_escape(noted_by if noted_by else " "), sign_name),
            ],
            [
                Paragraph(xml_escape(prepared_by_designation), sign_desig),
                Paragraph(xml_escape(noted_by_designation), sign_desig),
            ],
        ],
        colWidths=[doc.width / 2.0, doc.width / 2.0],
    )
    sign_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D0D0")),
                ("LINEBELOW", (0, 1), (0, 1), 0.8, colors.black),
                ("LINEBELOW", (1, 1), (1, 1), 0.8, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(sign_table)

    doc.build(story)
    return tmp.name
